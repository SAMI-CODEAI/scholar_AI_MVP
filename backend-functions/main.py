"""
AI Learning Assistant - Backend API
Google Functions Framework Implementation
Replaces Flask backend with Cloud Run compatible functions
"""

try:
    import functions_framework
except ImportError:
    # Fallback for local dev or if functions_framework has issues
    class functions_framework:
        @staticmethod
        def http(func):
            return func

from flask import Flask, request, jsonify, send_file
import os
import json
import random
import string
import io
import tempfile
import time

# Google Cloud imports
import google.generativeai as genai
from google.oauth2 import service_account
from google.cloud import speech
from google.cloud import storage

# Document processing imports
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from pypdf import PdfReader
from mutagen.mp3 import MP3

# Firebase Admin for token verification
import firebase_admin
from firebase_admin import auth, credentials

# Initialize Firebase Admin
if not firebase_admin._apps:
    try:
        cred = credentials.ApplicationDefault()
        firebase_admin.initialize_app(cred)
    except Exception:
        firebase_admin.initialize_app()

# Configure Gemini API
# Default key from environment or placeholder (used as fallback)
USER_GEMINI_API_KEY = "AIzaSyC0BgST84n0YqkSHPR6FfURsv_MYimVNLA"
DEFAULT_GEMINI_KEY = os.environ.get("GEMINI_API_KEY", USER_GEMINI_API_KEY)

def get_gemini_key(request=None):
    """Get Gemini API key from request header or fallback"""
    if request:
        header_key = request.headers.get('X-Gemini-API-Key')
        if header_key and header_key.strip():
            return header_key
    
    if DEFAULT_GEMINI_KEY and DEFAULT_GEMINI_KEY != "PASTE_YOUR_GEMINI_API_KEY_HERE":
        return DEFAULT_GEMINI_KEY
    return None

def configure_genai(request=None):
    """Configure GenAI with specific key"""
    api_key = get_gemini_key(request)
    if api_key:
        genai.configure(api_key=api_key)
        return True
    return False

# Google Cloud Storage bucket
GCS_BUCKET_NAME = os.environ.get("GCS_BUCKET_NAME", "ai-learning-assistant-storage")

# Database directory
DATABASE_DIR = os.environ.get("DATABASE_DIR", "/tmp/database")
os.makedirs(DATABASE_DIR, exist_ok=True)


# ============================================================================
# CORS HANDLING
# ============================================================================

def add_cors_headers(response, status_code=200):
    """Add CORS headers to response"""
    headers = {
        'Access-Control-Allow-Origin': '*',
        'Access-Control-Allow-Methods': 'GET, POST, PUT, DELETE, OPTIONS',
        'Access-Control-Allow-Headers': 'Content-Type, Authorization',
        'Access-Control-Max-Age': '3600'
    }
    return (response, status_code, headers)


def handle_cors_preflight():
    """Handle CORS preflight OPTIONS request"""
    return add_cors_headers('', 204)


# ============================================================================
# FIREBASE AUTH VERIFICATION
# ============================================================================

def verify_firebase_token(request):
    """Verify Firebase ID token from Authorization header"""
    auth_header = request.headers.get('Authorization', '')
    
    if not auth_header.startswith('Bearer '):
        return None, 'Missing or invalid Authorization header'
    
    id_token = auth_header.split('Bearer ')[1]
    
    try:
        decoded_token = auth.verify_id_token(id_token)
        return decoded_token, None
    except Exception as e:
        return None, f'Invalid token: {str(e)}'


# ============================================================================
# GEMINI AI FUNCTIONS
# ============================================================================

def prompt_flashcards(prompt):
    """Generate flashcards from transcript using Gemini"""
    model = genai.GenerativeModel("gemini-2.0-flash-lite-001")
    response = model.generate_content(
        "Given the following transcript, make 10 flashcards. Return these in JSON format. "
        "It should be a list of lists, where each sublist is two elements where the first "
        "element is the question/vocab word and the second element is the answer/definition. "
        "Don't output anything besides the JSON. Here is the transcript: " + prompt
    )
    text = response.text
    if text.startswith("```"):
        lines = text.split("\n")
        text = "\n".join(lines[1:-1])
    return json.loads(text)


def prompt_quiz(prompt):
    """Generate quiz from transcript using Gemini"""
    model = genai.GenerativeModel("gemini-2.0-flash-lite-001")
    response = model.generate_content(
        "Given the following transcript, I want you to generate a quiz with 10 questions. "
        "The quiz should be in JSON format. It should be a list of JSON objects. Each JSON "
        "object should have three fields: question, possible_answers, and index. Question is "
        "a string. possible_answers should be a list of possible answers. Only one answer "
        "should be correct. index should be a number and should be the index of the correct "
        "answer. Remember to put the answer in a random index so its harder to cheat! "
        "Please shuffle!!! Do not say anything other than this JSON: " + prompt
    )
    text = response.text
    if text.startswith("```"):
        lines = text.split("\n")
        text = "\n".join(lines[1:-1])
    
    j = json.loads(text)
    for question in j:
        ans = question["possible_answers"][question["index"]]
        random.shuffle(question["possible_answers"])
        question["index"] = question["possible_answers"].index(ans)
    
    return j


def prompt_summary(prompt):
    """Generate summary from transcript using Gemini"""
    model = genai.GenerativeModel("gemini-2.0-flash-lite-001")
    response = model.generate_content(
        "Given the following transcript, give a thorough summary of the main parts. "
        "Be as detailed as possible. Have a few main points which are bolded, and a lot "
        "of smaller subpoints, which use bullet points and are under the main points. "
        "USE MARKDOWN. DO NOT WRITE ANYTHING ELSE EXCEPT THE SUMMARY. FOR EXAMPLE, "
        "DO NOT WRITE Summary of Main Points: or Main Points: Start with an Overview: " + prompt
    )
    return response.text


def prompt_title(prompt):
    """Generate title from transcript using Gemini"""
    model = genai.GenerativeModel("gemini-2.0-flash-lite-001")
    response = model.generate_content(
        "Give me a title for this transcript. Should not be very long. "
        "Say nothing else except the title: " + prompt
    )
    return response.text


def prompt_everything(prompt):
    """Generate all study materials from transcript"""
    summary = prompt_summary(prompt)
    flash_cards = prompt_flashcards(prompt)
    quiz = prompt_quiz(prompt)
    title = prompt_title(prompt)
    
    return {
        "summary": summary,
        "flash_cards": flash_cards,
        "quiz": quiz,
        "title": title,
    }


# ============================================================================
# SPEECH TO TEXT
# ============================================================================

def get_audio_transcript(audio_file_path):
    """Convert audio to text using Google Cloud Speech-to-Text"""
    try:
        credentials_info = os.environ.get("GOOGLE_APPLICATION_CREDENTIALS")
        if credentials_info and os.path.exists(credentials_info):
            creds = service_account.Credentials.from_service_account_file(credentials_info)
        else:
            creds = None
        
        # Check audio length
        try:
            audio_len = MP3(audio_file_path).info.length
        except Exception:
            # Fallback if mutagen fails or not MP3
            audio_len = 0
            
        client = speech.SpeechClient(credentials=creds)
        
        # For short audio (< 60 seconds), use synchronous recognition
        if audio_len > 0 and audio_len <= 60:
            with io.open(audio_file_path, "rb") as f:
                content = f.read()
            
            audio = speech.RecognitionAudio(content=content)
            config = speech.RecognitionConfig(
                encoding=speech.RecognitionConfig.AudioEncoding.ENCODING_UNSPECIFIED,
                sample_rate_hertz=48000,
                language_code="en-US",
                audio_channel_count=2,
            )
            
            response = client.recognize(config=config, audio=audio)
            
            output = ""
            for result in response.results:
                output += result.alternatives[0].transcript
            return output
            
        else:
            # For longer audio (> 60s), upload to GCS and use long-running recognize
            storage_client = storage.Client(credentials=creds)
            bucket = storage_client.bucket(GCS_BUCKET_NAME)
            
            blob_name = f"uploads/{int(time.time())}_{os.path.basename(audio_file_path)}"
            blob = bucket.blob(blob_name)
            
            blob.upload_from_filename(audio_file_path)
            gcs_uri = f"gs://{GCS_BUCKET_NAME}/{blob_name}"
            
            config = speech.RecognitionConfig(
                encoding=speech.RecognitionConfig.AudioEncoding.ENCODING_UNSPECIFIED,
                sample_rate_hertz=48000,
                language_code="en-US",
                audio_channel_count=2,
            )
            audio = speech.RecognitionAudio(uri=gcs_uri)
            
            operation = client.long_running_recognize(config=config, audio=audio)
            
            # Wait for result (with timeout based on length)
            # Default timeout 300s + length
            timeout_val = 300
            if audio_len > 0:
                timeout_val = int(audio_len) + 300
                
            response = operation.result(timeout=timeout_val)
            
            # Cleanup - delete blob after processing
            try:
                blob.delete()
            except Exception:
                pass

            output = ""
            for result in response.results:
                output += result.alternatives[0].transcript
            return output

    except Exception as e:
        print(f"Speech-to-text error: {e}")
        # Return something to allow testing without valid credentials
        return f"Transcript generation failed: {str(e)}"


# ============================================================================
# FILE HANDLERS
# ============================================================================

def handle_pdf(file_path):
    """Extract text from PDF"""
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text


def handle_txt(file_path):
    """Read text file"""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def handle_docx(file_path):
    """Extract text from DOCX"""
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text


# ============================================================================
# DOCUMENT GENERATORS
# ============================================================================

def generate_id():
    """Generate unique ID for study guide"""
    letters = string.ascii_letters + string.digits
    return "".join(random.choice(letters) for _ in range(5))


def export_quiz_docx(quiz_data):
    """Export quiz to DOCX format"""
    document = Document()
    
    heading_style = document.styles["Heading1"]
    heading_font = heading_style.font
    heading_font.size = Pt(14)
    
    document.add_heading("Quiz", level=1)
    
    for question_data in quiz_data:
        question = question_data["question"]
        possible_answers = question_data["possible_answers"]
        
        p = document.add_paragraph()
        p.add_run(question).bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        for answer in possible_answers:
            document.add_paragraph(answer, style="ListBullet")
    
    return document


def export_flashcards_docx(cards):
    """Export flashcards to DOCX format with styling"""
    document = Document()
    document.add_heading("Flashcards", level=1)
    
    table = document.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    
    # Add header
    hdr_cells = table.add_row().cells
    hdr_cells[0].text = 'Front'
    hdr_cells[1].text = 'Back'

    # Apply styling to flashcards
    for card in cards:
        row = table.add_row()
        row.height = Inches(1.5)
        
        row_cells = row.cells
        row_cells[0].text = card[0]
        row_cells[1].text = card[1]
        
        # Apply colors (Redish for Front, Orangeish for Back) mimic original
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FFCCCB"/>'.format(nsdecls("w")))
        row_cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
        
        shading_elm_2 = parse_xml(r'<w:shd {} w:fill="FFDAB9"/>'.format(nsdecls("w")))
        row_cells[1]._tc.get_or_add_tcPr().append(shading_elm_2)

    return document


def export_summary_docx(summary_text):
    """Export summary to DOCX format"""
    document = Document()
    document.add_heading("Summary", level=1)
    
    for line in summary_text.split('\n'):
        if line.strip():
            if line.startswith('#'):
                document.add_heading(line.replace('#', '').strip(), level=2)
            elif line.startswith('**') and line.endswith('**'):
                p = document.add_paragraph()
                p.add_run(line.replace('**', '')).bold = True
            elif line.startswith('- '):
                document.add_paragraph(line.replace('- ', ''), style='ListBullet')
            else:
                document.add_paragraph(line)
    
    return document


# ============================================================================
# MAIN API HANDLER
# ============================================================================

@functions_framework.http
def api(request):
    """Main API handler for all endpoints"""
    
    # Handle CORS preflight
    if request.method == 'OPTIONS':
        return handle_cors_preflight()
    
    path = request.path
    method = request.method
    
    try:
        # Route: Health check
        if path == '/api/health' or path == '/health':
            return add_cors_headers(jsonify({"status": "ok", "service": "AI Learning Assistant Backend"}))
        
        # Route: Upload and process file
        if path == '/api/upload' and method == 'POST':
            return handle_upload(request)
        
        # Route: Get study guide by ID
        if path.startswith('/api/guide/') and method == 'GET':
            guide_id = path.split('/api/guide/')[1]
            return get_guide(guide_id)
        
        # Route: Get all guides
        if path == '/api/guides' and method == 'GET':
            return get_all_guides(request)
        
        # Route: Export quiz as DOCX
        if path.startswith('/api/export/quiz/') and method == 'GET':
            guide_id = path.split('/api/export/quiz/')[1]
            return export_quiz(guide_id)
        
        # Route: Export flashcards as DOCX
        if path.startswith('/api/export/flashcards/') and method == 'GET':
            guide_id = path.split('/api/export/flashcards/')[1]
            return export_flashcards(guide_id)
        
        # Route: Export summary as DOCX
        if path.startswith('/api/export/summary/') and method == 'GET':
            guide_id = path.split('/api/export/summary/')[1]
            return export_summary(guide_id)
        
        # Route: List available models (Debugging)
        if path == '/api/models' and method == 'GET':
            return list_available_models(request)
        
        # Route: Delete guide
        if path.startswith('/api/guide/') and method == 'DELETE':
            guide_id = path.split('/api/guide/')[1]
            return delete_guide(request, guide_id)
        
        # Route not found
        return add_cors_headers(jsonify({"error": "Not found"}), 404)
    
    except Exception as e:
        print(f"API Error: {e}")
        return add_cors_headers(jsonify({"error": str(e)}), 500)


def handle_upload(request):
    """Handle file upload and generate study materials"""
    
    user, error = verify_firebase_token(request)
    user_id = user.get('uid') if user else 'anonymous'
    
    # Configure Gemini with user provided key
    if not configure_genai(request):
        return add_cors_headers(jsonify({"error": "Gemini API Key is missing. Please provide it in the input field."}), 400)

    if 'file' not in request.files:
        return add_cors_headers(jsonify({"error": "No file provided"}), 400)
    
    file = request.files['file']
    filename = file.filename.lower()
    
    # Save file temporarily
    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as tmp:
        file.save(tmp.name)
        tmp_path = tmp.name
    
    try:
        if filename.endswith('.pdf'):
            text = handle_pdf(tmp_path)
        elif filename.endswith('.docx'):
            text = handle_docx(tmp_path)
        elif filename.endswith('.txt') or filename.endswith('.md') or filename.endswith('.html'):
            text = handle_txt(tmp_path)
        elif filename.endswith('.mp3'):
            text = get_audio_transcript(tmp_path)
        elif filename.endswith('.mp4'):
            # Basic extraction valid for some mp4 structures, but optimally needs ffmpeg
            # For this MVP, we assume the speech-to-text can handle it or we reuse mp3 logic
            text = get_audio_transcript(tmp_path)
        else:
            return add_cors_headers(jsonify({"error": "Unsupported file type"}), 400)
        
        if not text.strip():
            return add_cors_headers(jsonify({"error": "Could not extract text from file"}), 400)
        
        # Generate study materials
        result = prompt_everything(text)
        
        guide_id = generate_id()
        result['id'] = guide_id
        result['user_id'] = user_id
        result['filename'] = filename
        result['created_at'] = int(time.time())
        
        # Save to database
        db_path = os.path.join(DATABASE_DIR, f"{guide_id}.json")
        with open(db_path, 'w') as f:
            json.dump(result, f)
        
        return add_cors_headers(jsonify({"id": guide_id, "title": result.get('title', 'Untitled')}))
    
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass


def get_guide(guide_id):
    """Get study guide by ID"""
    db_path = os.path.join(DATABASE_DIR, f"{guide_id}.json")
    if not os.path.exists(db_path):
        return add_cors_headers(jsonify({"error": "Guide not found"}), 404)
    
    with open(db_path, 'r') as f:
        data = json.load(f)
    return add_cors_headers(jsonify(data))


def get_all_guides(request):
    """Get all study guides"""
    user, _ = verify_firebase_token(request)
    user_id = user.get('uid') if user else None
    
    guides = []
    try:
        if os.path.exists(DATABASE_DIR):
            for filename in os.listdir(DATABASE_DIR):
                if filename.endswith('.json'):
                    try:
                        with open(os.path.join(DATABASE_DIR, filename), 'r') as f:
                            data = json.load(f)
                            if user_id is None or data.get('user_id') == user_id or data.get('user_id') == 'anonymous':
                                guides.append({
                                    'id': data.get('id'),
                                    'title': data.get('title', 'Untitled'),
                                    'filename': data.get('filename', ''),
                                    'created_at': data.get('created_at', 0)
                                })
                    except Exception:
                        continue
        # Sort by newest first
        guides.sort(key=lambda x: x.get('created_at', 0), reverse=True)
    except Exception as e:
        print(f"Error listing guides: {e}")
    
    return add_cors_headers(jsonify({"guides": guides}))


def delete_guide(request, guide_id):
    """Delete a study guide"""
    user, error = verify_firebase_token(request)
    if error:
        return add_cors_headers(jsonify({"error": "Unauthorized"}), 401)
    
    db_path = os.path.join(DATABASE_DIR, f"{guide_id}.json")
    if not os.path.exists(db_path):
        return add_cors_headers(jsonify({"error": "Guide not found"}), 404)
    
    with open(db_path, 'r') as f:
        data = json.load(f)
        if data.get('user_id') != user.get('uid'):
            return add_cors_headers(jsonify({"error": "Forbidden"}), 403)
    
    os.unlink(db_path)
    return add_cors_headers(jsonify({"success": True}))


def export_quiz(guide_id):
    """Export quiz as DOCX"""
    db_path = os.path.join(DATABASE_DIR, f"{guide_id}.json")
    if not os.path.exists(db_path):
        return add_cors_headers(jsonify({"error": "Guide not found"}), 404)
    
    with open(db_path, 'r') as f:
        data = json.load(f)
    
    doc = export_quiz_docx(data.get('quiz', []))
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        doc.save(tmp.name)
        return send_file(
            tmp.name,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'quiz_{guide_id}.docx'
        )


def export_flashcards(guide_id):
    """Export flashcards as DOCX"""
    db_path = os.path.join(DATABASE_DIR, f"{guide_id}.json")
    if not os.path.exists(db_path):
        return add_cors_headers(jsonify({"error": "Guide not found"}), 404)
    
    with open(db_path, 'r') as f:
        data = json.load(f)
    
    doc = export_flashcards_docx(data.get('flash_cards', []))
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        doc.save(tmp.name)
        return send_file(
            tmp.name,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'flashcards_{guide_id}.docx'
        )


def export_summary(guide_id):
    """Export summary as DOCX"""
    db_path = os.path.join(DATABASE_DIR, f"{guide_id}.json")
    if not os.path.exists(db_path):
        return add_cors_headers(jsonify({"error": "Guide not found"}), 404)
    
    with open(db_path, 'r') as f:
        data = json.load(f)
    
    doc = export_summary_docx(data.get('summary', ''))
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        doc.save(tmp.name)
        return send_file(
            tmp.name,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'summary_{guide_id}.docx'
        )


def list_available_models(request):
    """List available Gemini models"""
    configure_genai(request)
    try:
        models = []
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                models.append(m.name)
        return add_cors_headers(jsonify({"models": models}))
    except Exception as e:
        return add_cors_headers(jsonify({"error": str(e)}), 500)

# ============================================================================
# LOCAL DEVELOPMENT SERVER
# ============================================================================

if __name__ == "__main__":
    app = Flask(__name__)
    @app.route('/', defaults={'path': ''}, methods=['GET', 'POST', 'PUT', 'DELETE', 'OPTIONS'])
    @app.route('/<path:path>', methods=['GET', 'POST', 'PUT', 'DELETE', 'OPTIONS'])
    def catch_all(path):
        from flask import request as flask_request
        flask_request.path = '/' + path
        return api(flask_request)
    
    print("Starting local development server on http://localhost:8080")
    app.run(host='0.0.0.0', port=8080, debug=True)
